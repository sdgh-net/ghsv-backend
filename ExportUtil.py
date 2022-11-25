import os
from docx import Document
from docx.oxml.ns import qn
from Config import Config, Const
from Functions import rm
from Supervision import Supervision
from Idea import Idea
from User import User

# REFerences
# https://zhuanlan.zhihu.com/p/375337204
# https://zhuanlan.zhihu.com/p/82880510
# https://blog.csdn.net/qq_40272386/article/details/114867630

DOCX_PREFIX = "docx导出_"
if os.path.isdir(Config.PATHS.TEMP_DIR):
    for i in os.listdir(Config.PATHS.TEMP_DIR):
        if i.startswith(DOCX_PREFIX):
            rm(os.path.join(Config.PATHS.TEMP_DIR, i))
Exported = []

def exportSupervisionToDocx(docx_file: str, sv: Supervision) -> None:
    global Exported
    if len(Exported)>100:
        for i in Exported[:50]:
            rm(i)
            Exported.remove(i)
    document = Document()
    document.styles['Normal'].font.name = Config.DOCX.FONT_NORMAL
    document.styles['Normal']._element.rPr.rFonts.set(
        qn('w:eastAsia'), Config.DOCX.FONT_NORMAL)
    heading = document.add_heading(level=0)
    heading_run = heading.add_run(f'{sv.getName()}反馈')
    heading_run.font.name = Config.DOCX.FONT_HEADING
    heading_run.element.rPr.rFonts.set(
        qn('w:eastAsia'), Config.DOCX.FONT_HEADING)

    ideas: list[Idea] = [i for i in sv.listIdeas() if i.getStatus() in (
        Idea.Const.IDEA_UNPROCESSED, Idea.Const.IDEA_ACCEPTED, Idea.Const.IDEA_REJECTED)]
    apartments: list[User] = User.listApartments()
    for u in apartments:
        my_ideas = [i for i in ideas if i.getTarget() == u.getUid()]
        if len(my_ideas) == 0:
            continue
        heading1 = document.add_heading(level=1)
        heading_run = heading1.add_run(u.getNickname())
        heading_run.font.name = Config.DOCX.FONT_APARTMENT_HEADING
        heading_run.element.rPr.rFonts.set(
            qn('w:eastAsia'), Config.DOCX.FONT_APARTMENT_HEADING)
        table = document.add_table(
            0, 2, style="Table Grid")
        for i in my_ideas:
            row_cells = table.add_row().cells
            row_cells[0].text = i.getContent()
            row_cells[1].text = i.toPrettyStr()
        table.autofit = True  # 没有什么用
    document.save(docx_file)
